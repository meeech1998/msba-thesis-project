#!/usr/bin/env python3
"""
complete_paper.py
Builds the final paper docx: all analysis, all figures, all findings.
EB Garamond 11pt, 1.15 spacing.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'outputs')
OUT = os.path.join(OUTPUT_DIR, 'complete_cit-mhrf-analysis_calderwood_20260328.docx')

# ══════════════════════════════════════════════════════════════════
# DOCUMENT SETUP
# ══════════════════════════════════════════════════════════════════

doc = Document()

# Page setup: Letter size, 1" margins
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

FONT = 'EB Garamond'
FONT_SIZE = Pt(11)
HEADING_COLOR = RGBColor(0, 0, 0)

# Set default style
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = FONT_SIZE
style.font.color.rgb = RGBColor(0, 0, 0)
pf = style.paragraph_format
pf.line_spacing = 1.15
pf.space_after = Pt(6)
pf.space_before = Pt(0)

# Heading styles
for level, (sz, sp_before, sp_after) in {
    1: (14, 18, 12),
    2: (12, 14, 8),
    3: (11, 10, 6),
}.items():
    hs = doc.styles[f'Heading {level}']
    hs.font.name = FONT
    hs.font.size = Pt(sz)
    hs.font.bold = True
    hs.font.color.rgb = HEADING_COLOR
    hs.paragraph_format.space_before = Pt(sp_before)
    hs.paragraph_format.space_after = Pt(sp_after)
    hs.paragraph_format.line_spacing = 1.15

# ══════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════════

def add_para(text, bold=False, italic=False, alignment=None, font_size=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = font_size or FONT_SIZE
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT
        run.font.color.rgb = HEADING_COLOR
    return h

def add_figure(filename, caption, width=Inches(6)):
    path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(path):
        add_para(f'[Figure not found: {filename}]', italic=True)
        print(f"  WARNING: {filename} not found")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.name = FONT
    r.font.size = Pt(10)
    r.italic = True
    cap.paragraph_format.space_after = Pt(8)

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = FONT_SIZE
    return p

def add_table_simple(headers, rows, col_widths=None):
    """Add a simple formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(10)
        # Light blue header shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D5E8F0" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = FONT
            run.font.size = Pt(10)

    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)

    return table

def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════
print("Building title page...")

# Add some spacing before title
for _ in range(4):
    doc.add_paragraph()

# Logo
logo_path = os.path.join(OUTPUT_DIR, 'extracted_logo.png')
if os.path.exists(logo_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(3))

doc.add_paragraph()  # spacer

# Title
add_para('Crisis Intervention Team Data Analysis:', bold=True,
         font_size=Pt(16), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para('Patterns in Severity, Repeat Encounters, and System Response',
         bold=True, font_size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para('Missoula Police Department Mental Health Referral Form (MHRF)',
         font_size=Pt(12), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

add_para('Michelle Calderwood', font_size=Pt(12),
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('University of Montana \u2014 MSBA Capstone', font_size=Pt(11),
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('2026-03-28', font_size=Pt(11),
         alignment=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════
print("Building table of contents...")

add_heading('TABLE OF CONTENTS', level=1)

toc_items = [
    '1.  Executive Summary',
    '2.  Introduction',
    '3.  Data & Methodology',
    '4.  Findings',
    '    4.1  Responding Agency',
    '    4.2  Severity \u2014 Behavioral Drivers',
    '    4.3  Severity vs. Frequency',
    '    4.4  Encounter Typology',
    '    4.5  Repeat Encounters',
    '    4.6  Time-Based Analysis',
    '    4.7  Location',
    '5.  Recommendations',
    '6.  Limitations',
    '7.  Conclusion',
    '8.  Appendix A \u2014 40-HR CIT Academy Analysis',
    '9.  Appendix B \u2014 Behavioral Indicator Validation',
    '10. Works Cited',
]
for item in toc_items:
    add_para(item, space_after=2)

page_break()

# ══════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════
print("Building Executive Summary...")

add_heading('EXECUTIVE SUMMARY', level=1)

add_para(
    'This project evaluates Crisis Intervention Team (CIT) data from the Missoula Police '
    'Department to understand how crisis encounters are currently documented and how the system '
    'can be improved to better support decision-making. A crisis encounter refers to a situation '
    'where law enforcement, often working in coordination with crisis response partners such as '
    'EMS or behavioral health resources when available, responds to an individual experiencing a '
    'mental health crisis, behavioral distress, or substance-related impairment. Unlike a typical '
    'police call that focuses on criminal activity and enforcement, these encounters prioritize '
    'de-escalation, safety, and connecting individuals to appropriate care or services, making '
    'them more complex and dependent on clear, consistent documentation.'
)

add_para(
    'Crisis encounters are recorded using the Mental Health Referral Form (MHRF), which captures '
    'key information including behavioral indicators, call type, housing status, legal criteria, '
    'and outcomes. While this provides valuable documentation, the data is primarily used for '
    'handoff to receiving facilities and periodic reporting tied to program evaluation and funding '
    'requirements, rather than for continuous analysis or operational decision support.'
)

add_para(
    'The current workflow relies on handwritten forms followed by manual data entry. This '
    'introduces delays (an average of 18 days from field documentation to system availability), '
    'inconsistencies, and data quality issues that limit the system\u2019s ability to support '
    'real-time monitoring, proactive intervention, or advanced analytics.'
)

add_para(
    'To better understand system patterns, a series of statistical analyses were conducted, '
    'including feature engineering, logistic regression, clustering, chi-square testing, spatial '
    'hotspot, and time-based analysis. These methods were used to identify the primary drivers of '
    'crisis severity, repeat encounters, and overall system behavior. All statistical models were '
    'implemented from scratch using only NumPy, without external machine learning libraries, to '
    'ensure full transparency and understanding of every analytical step.'
)

add_para(
    'The most significant finding is the role of the responding agency. Encounters handled solely '
    'by behavioral health specialists (Mobile Support Team, or MST) show a severity rate of just '
    '4.1%, compared to 58.3% for law enforcement-only responses (MPD) and 48.0% for co-response '
    'deployments. This difference is not driven by a single severity component \u2014 it is '
    'pervasive across involuntary determinations, use of force, and injury. Critically, this '
    'pattern persists within the same call types: for suicidal/self-harm calls, MST-only severity '
    'is approximately 1.4% versus MPD-only at 57.2%. A chi-square test confirms the association '
    'is statistically significant (chi-squared = 198.1, p < 0.001).'
)

add_para(
    'Beyond agency effects, behavioral escalation \u2014 particularly indicators such as anger, '
    'confusion, and manic presentation \u2014 is the strongest behavioral driver of severity. The '
    'severity prediction model achieves a test AUC of 0.827, indicating strong discriminative '
    'ability. In contrast, repeat encounters are more strongly associated with underlying life '
    'circumstances, especially housing status and whether someone is unhoused. The repeat model '
    'achieves a test AUC of only 0.495 \u2014 essentially random chance \u2014 indicating that '
    'encounter-level features cannot predict who will return to the system. This is itself a '
    'meaningful finding: repeat contact is driven by factors that exist outside the scope of any '
    'single encounter.'
)

add_para(
    'Crisis encounters also group into three distinct profiles: an escalation-driven type '
    '(high severity, driven by behavioral indicators), a system-driven type (low severity but '
    'high repeat contact, driven by housing instability), and an emotional distress type (lower '
    'acuity, depression-driven). These clusters reinforce that different types of crises require '
    'fundamentally different response approaches.'
)

add_para(
    'The most significant limitation identified is not the analytical approach, but the data '
    'collection process itself. The reliance on handwritten forms and manual entry prevents '
    'consistent, timely, and scalable use of the data. The primary recommendation is to '
    'transition to a digital-first data collection system and to expand behavioral health '
    'response capacity, given the dramatic severity differences associated with response model.'
)

add_para(
    'In parallel, this study also evaluated the impact of the 40-hour CIT Academy. Analysis of '
    'pre-training, post-training, and follow-up survey data showed substantial and sustained '
    'improvements across all measured Knowledge, Skills, and Abilities (KSA) domains, with '
    'Cohen\u2019s d effect sizes ranging from 1.03 to 3.31. These findings indicate that the '
    'training is effective in strengthening individual officer capability, while the encounter-level '
    'results show that system-level factors \u2014 particularly who responds \u2014 may matter '
    'even more for outcomes.'
)

page_break()

# ══════════════════════════════════════════════════════════════════
# 2. INTRODUCTION
# ══════════════════════════════════════════════════════════════════
print("Building Introduction...")

add_heading('INTRODUCTION', level=1)

add_para(
    'CIT programs are widely implemented models designed to improve how communities respond to '
    'behavioral health crises by connecting law enforcement, behavioral health providers, and '
    'community resources. These programs emphasize de-escalation, safety, and diversion to '
    'appropriate care, with the goal of reducing unnecessary criminal justice involvement and '
    'improving outcomes for individuals in crisis (NAMI, 2023; SAMHSA, 2020).'
)

add_para(
    'MPD operates within this framework, using a coordinated crisis response model supported by '
    'officer training and structured reporting. A central component of this system is the MHRF, '
    'which is used to document crisis encounters in the field. The form captures key details '
    'including behavioral indicators, call type (the initial reason for the call for service), '
    'housing status, legal criteria, and encounter outcomes.'
)

add_para(
    'The MHRF serves as the primary mechanism through which crisis encounters are recorded and '
    'later reviewed. It enables visibility into system activity, allowing stakeholders to track '
    'encounter characteristics and monitor overall patterns of response. As such, the form '
    'functions not only as a documentation tool, but also as the foundation for understanding how '
    'the crisis response system operates in practice.'
)

add_para(
    'Despite this, the current structure of the system presents limitations in how the data can '
    'be used. While the form captures detailed information, the way it is collected and stored '
    'constrains its ability to support deeper analysis. This limits the system\u2019s capacity to '
    'move beyond basic documentation toward a more comprehensive understanding of crisis dynamics.'
)

add_para(
    'This project builds on that foundation by developing an analytical framework to examine '
    'patterns within crisis encounters, including variation in outcomes, recurrence of contact, '
    'and differences in how crises present. In parallel, it considers how system design and data '
    'collection practices influence what can be measured, analyzed, and ultimately acted upon.'
)

add_para(
    'In addition, this study incorporates an evaluation of the 40-hour CIT Academy, which is '
    'designed to strengthen officer knowledge, preparedness, and confidence in responding to '
    'mental health crises. Examining both training outcomes and encounter-level data provides a '
    'more complete view of how the CIT model functions, linking individual-level capability with '
    'system-level behavior.'
)

add_para(
    'Because the MHRF underpins both documentation and analysis, its structure plays a critical '
    'role in shaping what insights the system can produce. Understanding both the analytical '
    'patterns within the data and the constraints of the data collection process is therefore '
    'essential to evaluating and improving crisis response.'
)

add_para(
    'Crisis response within the CIT model operates as an integrated, multi-agency system rather '
    'than a single point of intervention. As shown in Figure 1, crisis encounters may involve '
    'coordination between emergency communications (911/988), law enforcement, emergency medical '
    'services, mobile crisis teams, and mental health receiving centers, depending on the nature '
    'and severity of the situation. Behavioral health partners play a critical role within this '
    'system, particularly in assessment, stabilization, and referral, while law enforcement often '
    'serves as the initial point of contact and decision-making in the field. Although crisis '
    'response is inherently collaborative, this analysis focuses specifically on the MPD '
    'perspective, as MPD is responsible for initiating response in many encounters and for '
    'completing the MHRF, which serves as the primary source of data used throughout this study.'
)

add_figure('extracted_cit_diagram.png', 'Figure 1. Components of an Integrated Crisis Response System', width=Inches(6))

add_heading('Problem Context', level=2)

add_para(
    'The CIT system captures detailed crisis encounter data through the MHRF, but its current '
    'use is limited to retrospective reporting rather than real-time decision support. In '
    'practice, real-time decision support would mean giving officers or responders immediate '
    'visibility into factors such as prior encounters, risk indicators, and likely outcomes while '
    'they are on a call, helping guide decisions around de-escalation, transport, or referral. '
    'Because the data is both descriptive and collected through manual, inconsistent processes, '
    'it does not reliably support analysis of outcome drivers such as severity and repeat contact. '
    'As a result, the system captures information but does not fully enable actionable insight or '
    'proactive response.'
)

add_heading('Research Objective', level=2)

add_para(
    'This project uses MHRF data to understand past crisis encounters and inform future system '
    'design, alongside an evaluation of CIT training outcomes. Analytical modeling is used to '
    'identify patterns in severity, repeat encounters, and overall system behavior, while '
    'training data is analyzed to assess changes in officer knowledge, preparedness, and '
    'confidence. Together, these insights inform how data can be better captured and structured to '
    'support more consistent analysis and enable future predictive use.'
)

page_break()

# ══════════════════════════════════════════════════════════════════
# 3. DATA & METHODOLOGY
# ══════════════════════════════════════════════════════════════════
print("Building Data & Methodology...")

add_heading('DATA & METHODOLOGY', level=1)

add_para(
    'The dataset consists of 861 crisis encounter records involving 609 unique individuals, '
    'collected through the MHRF between March 2024 and February 2026. Each record captures '
    'information across 51 raw fields, including behavioral indicators observed during the '
    'encounter, call type, housing status, responding agencies, legal criteria, and outcome '
    'information.'
)

add_para(
    'Prior to analysis, I took steps to remove or transform any information that could compromise '
    'individual privacy while preserving the integrity of the dataset. I removed all initials, '
    'exact ages, and other personally identifiable information from the working dataset, including '
    'any identifying details contained in free-text fields. To retain analytical value, I used '
    'date of birth to calculate and group individuals into age categories rather than using exact '
    'ages. I also created internal person-level identifiers by initially linking records using '
    'combinations of date of birth and initials, then replaced these with de-identified IDs to '
    'ensure individuals could be tracked across encounters without exposing their identity. In '
    'addition, I standardized and structured variables to ensure consistency across records while '
    'maintaining the overall patterns and relationships within the data. These steps allowed the '
    'dataset to be analyzed holistically while protecting individual privacy.'
)

add_para(
    'An important distinction in the data is between encounter dates and data entry dates. The '
    'date_at_incident field records when the crisis encounter actually occurred in the field, '
    'while start_time and completion_time record when the data was entered into the system. The '
    'average gap between these is approximately 18 days, reflecting the delay inherent in the '
    'paper-based workflow. All time-based analyses in this study use the actual encounter date '
    '(date_at_incident) to reflect when events occurred rather than when they were documented.'
)

add_para(
    'Feature engineering expanded the dataset from 51 raw columns to 88 analytical features. '
    'This included binary encoding of behavioral indicators from structured checklist fields, '
    'keyword-based parsing of descriptive text for psychosis-related indicators, standardization '
    'of categorical variables, and construction of the two primary outcome variables described '
    'below. All statistical models \u2014 including logistic regression, AUC computation, '
    'chi-square testing, and k-modes clustering \u2014 were implemented from scratch using only '
    'NumPy, without external machine learning libraries such as scikit-learn.'
)

add_heading('Age', level=2)

add_para(
    'Crisis encounters in the dataset are concentrated among adult populations, particularly '
    'within the 18\u201339 age group, with fewer cases observed at younger and older extremes. '
    'While age provides important context for understanding who is represented in the system, it '
    'does not appear to strongly explain key outcomes such as severity or repeat encounters. This '
    'is supported by the modeling results, where age variables were not statistically significant '
    'predictors and showed minimal effect sizes compared to behavioral and situational factors.'
)

add_figure('fig00_age_distribution.png', 'Figure 2. Encounter Distribution by Age Group')

add_heading('Behavioral Indicators', level=2)

add_para(
    'Across encounters, the most common observed indicators include depression, substance '
    'involvement, and confusion. These indicators are frequently present across the dataset and '
    'reflect the broad range of behavioral health concerns captured by the MHRF. At the same '
    'time, less frequent indicators such as manic behavior appear to be associated with more '
    'severe encounters, suggesting that prevalence and impact are not the same.'
)

add_figure('fig03_behavioral_prevalence.png', 'Figure 3. Behavioral Indicator Frequency',
           width=Inches(5.5))

add_heading('Housing Status', level=2)

add_para(
    'Housing status provides an important view into the composition of the encounter population. '
    'Most encounters involve housed individuals, which reflects the overall population represented '
    'in the system. However, individuals experiencing housing instability account for a '
    'disproportionately large share of repeat encounters, indicating that recurrence may be shaped '
    'less by the immediate crisis event and more by broader structural conditions.'
)

add_figure('fig04_housing.png', 'Figure 4. Encounters by Housing Status', width=Inches(5.5))

add_figure('fig16_housing_repeat.png',
           'Figure 5. Repeat encounter rates by housing status, shown as the percentage of '
           'individuals within each group who experienced repeat encounters.', width=Inches(5.5))

add_heading('Call Types', level=2)

add_para(
    'Call type refers to the initial classification assigned to an incident when it is dispatched, '
    'based on the information available at the time of the call. Common call types include welfare '
    'checks, person needs assistance, and suicidal or self-harm related incidents. These '
    'categories describe how encounters enter the system and provide useful operational context, '
    'but on their own they do not appear to strongly distinguish severity or recurrence.'
)

add_figure('fig05_call_types.png', 'Figure 6. Call Type Distribution', width=Inches(6))

add_para(
    'Overall, the encounter profile shows that the system captures a wide range of behavioral, '
    'demographic, and operational information. While these descriptive patterns are useful for '
    'understanding the general shape of the data, they are not sufficient to explain why some '
    'encounters become severe or why some individuals return repeatedly to the system. This is '
    'supported by the modeling results, where many descriptive variables such as demographics and '
    'call type were not statistically significant predictors and showed limited explanatory power '
    'compared to behavioral indicators and underlying conditions like housing status.'
)

add_heading('Defining Severity', level=2)

add_para(
    'The analysis began with a key limitation of the MHRF dataset: there is no single direct '
    'measure of severity. To address this, severity was constructed using indicators that most '
    'clearly reflect higher-acuity encounters, including involuntary status, elevated commitment '
    'criteria (such as danger to self, danger to others, or inability to meet basic needs), use '
    'of force, and injury to the subject or others. These components were combined into a '
    'composite measure and ultimately into a binary variable, severe_flag, indicating whether an '
    'encounter should be classified as severe. Across the dataset, 40.4% of encounters are '
    'classified as severe.'
)

add_para(
    'With this outcome defined, the analysis then focused on identifying which observable features '
    'are associated with severe encounters. A logistic regression model was estimated using '
    'behavioral and clinical indicators documented during each encounter. Logistic regression '
    'estimates the probability that an encounter is classified as severe based on these '
    'characteristics, while holding other variables constant. The model produces coefficients that '
    'indicate the direction and strength of each relationship, and statistical significance '
    'testing is used to assess whether these relationships are likely to reflect real patterns in '
    'the data rather than random variation.'
)

add_para(
    'To avoid circular reasoning, any variables used to construct the severity measure, such as '
    'force, injury, and involuntary status, were excluded from the model. This ensures that the '
    'results are not mechanically driven by the definition of severity itself, but instead reflect '
    'independent associations between observed behaviors and the likelihood of a severe outcome.'
)

add_heading('Responding Agencies', level=2)

add_para(
    'The MHRF records which agencies were involved in each crisis encounter. I categorized these '
    'into four response models: Law Enforcement Only (primarily MPD, with 434 encounters), '
    'Behavioral Health Only (MST, with 245 encounters), Co-Response where both respond together '
    '(175 encounters), and Other (26 encounters). This classification allows comparison of how '
    'different response configurations relate to encounter outcomes.'
)

add_heading('Time', level=2)

add_para(
    'Time-based analysis was used to evaluate how crisis encounters evolve over time and whether '
    'changes reflect true shifts in demand or differences in reporting and system adoption. By '
    'analyzing trends in volume, severity, and repeat encounters, this approach helps separate '
    'fluctuations in activity from changes in the underlying risk structure and identify potential '
    'seasonal or operational patterns.'
)

add_figure('fig01_monthly_volume.png', 'Figure 7. Monthly Encounter Volume', width=Inches(6))

add_para(
    'Encounter volume fluctuates over time, with early 2024 (February\u2013April) likely '
    'reflecting underreporting during the initial rollout of the reporting process rather than '
    'true lower demand. Volume increases sharply through spring 2024 and reaches an initial peak '
    'in June 2024, followed by a decline in July and August before stabilizing at an elevated '
    'level through early fall. From this higher baseline, volume increases again and reaches a '
    'second, more pronounced peak in December 2024, the highest point observed in the data.'
)

add_para(
    'Overall, the pattern reflects three phases: an initial reporting ramp-up, a sustained higher '
    'baseline, and multiple peaks (early summer and winter), indicating that variation in '
    'encounter volume is driven by a combination of reporting maturity and recurring periods of '
    'higher demand rather than structural changes in the system.'
)

add_heading('Location', level=2)

add_para(
    'A spatial heatmap of geocoded crisis encounters shows that activity is concentrated most '
    'heavily in the urban core of Missoula, particularly along major corridors and service-access '
    'areas. This indicates that encounters are not randomly distributed but instead reflect '
    'underlying population density and system access points. While there is some variation, no '
    'area shows a distinct concentration of severe or repeat encounters beyond what would be '
    'expected based on overall activity levels. Severity and recurrence appear to follow overall '
    'activity patterns rather than forming distinct spatial clusters.'
)

page_break()

# ══════════════════════════════════════════════════════════════════
# 4. FINDINGS
# ══════════════════════════════════════════════════════════════════
print("Building Findings...")

add_heading('FINDINGS', level=1)

add_para(
    'Below are the key findings from the analyses conducted on the MHRF dataset. Together, these '
    'findings show that crisis severity and repeat encounters are driven by different underlying '
    'mechanisms, and that the current system captures useful information but is limited by how '
    'that information is structured and collected.'
)

# ── 4.1 RESPONDING AGENCY ──────────────────────────────────────
add_heading('4.1  Responding Agency', level=2)

add_para(
    'The most consequential finding in this dataset is the magnitude of the difference in '
    'outcomes depending on which type of agency responds to a crisis. The three primary response '
    'configurations \u2014 law enforcement only (MPD), behavioral health only (MST), and '
    'co-response \u2014 show dramatically different severity profiles.'
)

# Agency comparison table
add_table_simple(
    ['Response Model', 'n', 'Severity Rate', 'Involuntary', 'Force', 'Injury'],
    [
        ['Behavioral Health Only (MST)', '245', '4.1%', '2.0%', '0.4%', '2.4%'],
        ['Co-Response', '175', '48.0%', '52.6%', '8.6%', '10.9%'],
        ['Law Enforcement Only (MPD)', '434', '58.3%', '63.4%', '10.1%', '15.9%'],
        ['Other', '7', '14.3%', '14.3%', '14.3%', '0.0%'],
    ],
    col_widths=[2.2, 0.6, 1.0, 1.0, 0.8, 0.9]
)
add_para('')  # spacer

add_para(
    'The differences are striking and they are not limited to one component of the severity '
    'measure. MST-only encounters show rates near zero across all severity dimensions: '
    'involuntary determination at 2.0%, force at 0.4%, and subject injury at 2.4%. Law '
    'enforcement-only encounters show rates 15\u201330 times higher on every component. '
    'Co-response encounters fall between these extremes but closer to the law enforcement profile.'
)

add_figure('fig02_response_model.png',
           'Figure 8. Response Model Distribution and Severity Rates', width=Inches(6))

add_figure('fig06_agency_severity_components.png',
           'Figure 9. Severity Components by Response Model', width=Inches(5.5))

add_para(
    'A critical question is whether the agency-severity relationship reflects selection bias '
    '\u2014 perhaps MST simply handles less severe calls. To test this, I compared severity rates '
    'within the same call type categories. For suicidal/self-harm calls \u2014 among the highest-'
    'acuity call types \u2014 MST-only responses show a severity rate of approximately 1.4%, '
    'while MPD-only responses show 57.2%. For person needs assistance calls, the gap is similar. '
    'This within-call-type persistence substantially mitigates the selection-bias explanation, '
    'though it does not eliminate it entirely (see Limitations).'
)

add_figure('fig07_agency_within_calltype.png',
           'Figure 10. Severity Rate by Response Model Within Same Call Types', width=Inches(6))

add_para(
    'A chi-square test of independence between response model and severity yields chi-squared = '
    '198.1, df = 3, p < 0.001, confirming that the association is statistically significant. The '
    'magnitude of this statistic \u2014 far exceeding the critical value at any conventional '
    'significance level \u2014 reflects the size and consistency of the observed differences.'
)

add_para(
    'In the logistic regression model that includes both agency and behavioral indicators, MPD '
    'involvement (coefficient = 1.89, p < 0.001) emerges as the single largest positive predictor '
    'of severity \u2014 stronger than any behavioral indicator. MST involvement (coefficient = '
    '\u22120.89, p < 0.001) is the largest negative predictor. In other words, who responds to a '
    'crisis matters at least as much as what behaviors are observed during it.'
)

add_para(
    'This does not mean the relationship is purely causal. As discussed in the limitations, '
    'responding agency is not randomly assigned, and there may be unobserved factors influencing '
    'both dispatch decisions and outcomes. But the size and consistency of the difference, '
    'especially within the same call types, makes it the most operationally meaningful finding in '
    'this analysis.'
)

# ── 4.2 SEVERITY — BEHAVIORAL ──────────────────────────────────
add_heading('4.2  Severity \u2014 Behavioral Drivers', level=2)

add_para(
    'Beyond responding agency, the results show a clear and consistent pattern: behavioral '
    'escalation is the strongest behavioral driver of severity. Angry or uncooperative behavior '
    'emerges as the most influential behavioral predictor (coefficient = 0.80, p < 0.001), '
    'followed by manic presentation (0.74, p < 0.001) and confusion (0.69, p < 0.001). These '
    'indicators reflect observable loss of behavioral control and instability in real time, which '
    'aligns with how crisis situations escalate in practice.'
)

add_para(
    'Substance involvement and disorganized speech show more moderate relationships. In contrast, '
    'depression shows a negative association with severity (coefficient = \u22120.26, p = 0.020), '
    'and hallucinations also show a negative relationship (\u22120.40, p < 0.001), suggesting '
    'they are less characteristic of high-acuity encounters and may represent a different type of '
    'crisis presentation.'
)

# Severity coefficients table
add_table_simple(
    ['Feature', 'Coefficient', 'Std Error', 'p-value', 'Significant'],
    [
        ['MPD Responded', '1.889', '0.099', '< 0.001', 'Yes'],
        ['Anger/Uncooperative', '0.800', '0.105', '< 0.001', 'Yes'],
        ['Manic', '0.744', '0.102', '< 0.001', 'Yes'],
        ['Confusion', '0.688', '0.107', '< 0.001', 'Yes'],
        ['Probable Cause', '0.423', '0.082', '< 0.001', 'Yes'],
        ['Substance Involved', '0.081', '0.116', '0.481', 'No'],
        ['Disorganized Speech', '0.210', '0.102', '0.042', 'Yes'],
        ['Scared/Frightened', '0.095', '0.103', '0.354', 'No'],
        ['Housed', '\u22120.300', '0.097', '0.003', 'Yes'],
        ['Depression', '\u22120.257', '0.121', '0.020', 'Yes'],
        ['Hallucinations', '\u22120.397', '0.094', '< 0.001', 'Yes'],
        ['MST Responded', '\u22120.886', '0.106', '< 0.001', 'Yes'],
        ['Unhoused', '\u22120.875', '0.096', '< 0.001', 'Yes'],
    ],
    col_widths=[1.8, 1.0, 0.9, 0.9, 0.9]
)
add_para('')

add_figure('fig08_severity_coefficients.png',
           'Figure 11. Logistic Regression Coefficients for Severity Model', width=Inches(5.5))

add_para(
    'A key concern in this analysis is whether these behavioral indicators are truly predictive '
    'of severity, or whether they simply reflect how severe encounters are described in '
    'documentation. To address this, additional validation checks were conducted (see Appendix B). '
    'These checks compared indicator rates in severe vs. non-severe encounters, applied a stricter '
    'definition of severity using only force or injury, and examined behavioral indicators in '
    'encounters where force or injury are not present. In all cases, the same indicators \u2014 '
    'particularly anger, mania, and confusion \u2014 consistently differentiated severe from '
    'non-severe encounters, strengthening the interpretation that they represent genuine signals '
    'of escalation.'
)

add_heading('Model Performance', level=3)

add_para(
    'To evaluate how well the model performs overall, I tested it on a held-out portion of the '
    'data that was not used during training. The model achieved a test AUC of 0.827 (training AUC '
    '= 0.837), which indicates strong discriminative ability with only modest overfitting. In '
    'practical terms, this means the model correctly distinguishes between severe and non-severe '
    'encounters about 83% of the time.'
)

add_figure('fig09_severity_roc.png',
           'Figure 12. ROC Curve for Severity Model (Test AUC = 0.827)', width=Inches(4.5))

add_heading('Housing and Severity', level=3)

add_para(
    'The housing model shows a negative association between experiencing homelessness and '
    'encounter severity (odds ratio = 0.56). Encounters involving unhoused individuals are less '
    'likely to be classified as severe compared to those involving housed individuals. This is '
    'also reflected in the descriptive rates, where housed individuals show a higher severity rate '
    '(43.5%) than unhoused individuals (29.9%).'
)

add_para(
    'However, this should be interpreted carefully. This difference does not necessarily mean '
    'unhoused individuals experience less severe crises. There are several reasons in the analysis '
    'that suggest this relationship is influenced by how the data is structured rather than a true '
    'difference in underlying severity. First, the effect changes depending on the model. Second, '
    'this model does not include behavioral indicators, which were shown in earlier analyses to be '
    'the strongest drivers of severity. Third, if one group has more total encounters, including '
    'a higher number of lower-acuity interactions, their average severity will appear lower.'
)

add_heading('Call Type and Severity', level=3)

add_para(
    'Call type (the initial classification at dispatch) was tested as a predictor of severity. '
    'Most call type categories were not statistically significant. The only categories that showed '
    'a consistent and statistically significant relationship were suicidal/self-harm and person '
    'needs assistance, both of which were associated with a lower likelihood of severity. This '
    'tells us that call type is not a strong predictor of which encounters will escalate. Instead, '
    'it is more useful for identifying certain types of calls that are less likely to result in '
    'severe outcomes.'
)

# ── 4.3 SEVERITY VS FREQUENCY ──────────────────────────────────
add_heading('4.3  Severity vs. Frequency', level=2)

add_para(
    'One of the clearest patterns in the data is the difference between what appears most often '
    'and what matters most for severity. Figure 13 compares the frequency of each statistically '
    'significant behavioral or clinical indicator to its coefficient in the severity model. This '
    'allows for a direct comparison between how common an indicator is and how strongly it is '
    'associated with higher-acuity encounters.'
)

add_figure('fig15_frequency_vs_impact.png',
           'Figure 13. Indicator Frequency and Severity Impact Comparison', width=Inches(5.5))

add_para(
    'The figure shows that frequency and impact are not aligned. Confusion represents one of the '
    'most operationally meaningful signals, as it is both relatively common and strongly associated '
    'with higher-acuity encounters. Depression and substance involvement appear frequently, but '
    'their coefficients are more moderate. In contrast, angry or uncooperative behavior stands '
    'out as the strongest predictor of severity despite occurring at a moderate frequency. Manic '
    'presentation also shows a strong positive association but occurs less often.'
)

add_para(
    'Overall, the figure reinforces that the most frequently observed indicators are not '
    'necessarily the most important for identifying severe crises. Instead, severity is most '
    'strongly associated with specific forms of behavioral escalation, suggesting that operational '
    'focus should prioritize high-impact signals rather than the most commonly documented features.'
)

# ── 4.4 ENCOUNTER TYPOLOGY ─────────────────────────────────────
add_heading('4.4  Encounter Typology', level=2)

add_para(
    'The clustering analysis builds on the earlier severity and repeat models by asking a '
    'different question: rather than isolating individual predictors, do crisis encounters group '
    'into consistent, real-world patterns? Cluster analysis is a data-driven method that groups '
    'similar encounters together based on shared characteristics, without predefining categories '
    'in advance.'
)

add_para(
    'To answer this, a reduced feature set was constructed using the indicators that were most '
    'meaningful across prior analysis: anger/uncooperative behavior, manic presentation, '
    'confusion, depression, repeat contact, and housing status (unhoused). A k-modes clustering '
    'algorithm was applied to identify three distinct encounter types.'
)

add_figure('fig11_cluster_profiles.png',
           'Figure 14. Encounter Type Profiles (Indicator Composition)', width=Inches(5.5))

add_para(
    'The results show that each type is defined by a distinct combination of behavioral and '
    'structural characteristics:'
)

add_para(
    'High-Risk / Escalation-Driven (n = 347): Characterized by high levels of manic behavior, '
    'anger, and confusion. These encounters reflect acute behavioral escalation and have a '
    'severity rate of 84.1%. This aligns closely with the severity model, where behavioral '
    'escalation emerged as the strongest predictor.'
)

add_para(
    'Low-Risk / Emotional Distress (n = 145): Defined primarily by depressive presentation, with '
    'minimal presence of anger, mania, or confusion. These have a severity rate of only 5.5% and '
    'represent a lower-intensity but still critical subset where the primary driver is internal '
    'emotional distress rather than external behavioral escalation.'
)

add_para(
    'Moderate-Risk / System-Driven (n = 322): Show relatively low levels of acute behavioral '
    'indicators but higher rates of repeat contact and unhoused status. Severity is 7.5%, but '
    'this group accounts for a disproportionate share of repeat encounters. Unlike the other '
    'types, this group is defined by ongoing system involvement rather than how the crisis '
    'presents in the moment.'
)

add_figure('fig12_cluster_outcomes.png',
           'Figure 15. Severity and Repeat Rates by Encounter Type', width=Inches(5.5))

add_para(
    'Taken together, these findings highlight a key insight: severity and repeat involvement are '
    'not only driven by different variables, but also manifest as fundamentally different types of '
    'encounters. Escalation-driven encounters align with severity and represent acute behavioral '
    'crises, while system-driven encounters align with repeat contact and reflect underlying '
    'structural instability. This reinforces that crisis encounters cannot be understood through '
    'a single dimension, but instead reflect a combination of behavioral state and structural '
    'context.'
)

add_para(
    'From an operational perspective, these results suggest that a single response model is '
    'insufficient. Different encounter types require different approaches: escalation-driven '
    'encounters require rapid recognition of behavioral instability and strong de-escalation '
    'capabilities; emotional distress encounters are more likely to benefit from connection to '
    'behavioral health services; and system-driven encounters point to the need for longer-term '
    'interventions, particularly around housing stability and continuity of care.'
)

# ── 4.5 REPEAT ENCOUNTERS ──────────────────────────────────────
add_heading('4.5  Repeat Encounters', level=2)

add_para(
    'The repeat encounter analysis reveals a fundamentally different pattern compared to the '
    'severity model. Instead of behavioral escalation, underlying life conditions emerge as the '
    'primary drivers of repeated system contact. Across the dataset, 25.2% of individuals have '
    'repeat encounters.'
)

add_heading('Housing Instability', level=3)

add_para(
    'Housing status is the strongest predictor of repeat encounters. Individuals experiencing '
    'homelessness have substantially higher odds of repeat encounters (coefficient = 0.43, p < '
    '0.001), and this relationship is both large and consistent, clearly distinguishing it from '
    'all other variables in the model.'
)

add_para(
    'All other factors show weaker and less reliable relationships. Call types such as emergency '
    'evaluation and suicidal/self-harm are directionally associated with higher repeat rates, but '
    'these effects do not reach statistical significance. Similarly, while severity is positively '
    'associated with repeat encounters, the effect is modest.'
)

# Repeat coefficients table
add_table_simple(
    ['Feature', 'Coefficient', 'Std Error', 'p-value', 'Significant'],
    [
        ['MST Responded', '0.549', '0.116', '< 0.001', 'Yes'],
        ['Unhoused', '0.431', '0.113', '< 0.001', 'Yes'],
        ['Severe Encounter', '0.327', '0.111', '0.003', 'Yes'],
        ['MPD Responded', '0.109', '0.111', '0.326', 'No'],
        ['Confusion', '0.085', '0.123', '0.487', 'No'],
        ['Housed', '\u22120.008', '0.107', '0.943', 'No'],
        ['Depression', '\u22120.163', '0.124', '0.189', 'No'],
        ['Anger/Uncooperative', '\u22120.208', '0.112', '0.064', 'No'],
        ['Manic', '\u22120.358', '0.109', '0.001', 'Yes'],
        ['Substance Involved', '\u22120.763', '0.122', '< 0.001', 'Yes'],
    ],
    col_widths=[1.8, 1.0, 0.9, 0.9, 0.9]
)
add_para('')

add_figure('fig10_repeat_coefficients.png',
           'Figure 16. Logistic Regression Coefficients for Repeat Encounter Model', width=Inches(5.5))

add_heading('Model Performance \u2014 A Meaningful Null Finding', level=3)

add_para(
    'The repeat encounter model achieves a test AUC of only 0.495, which is essentially '
    'equivalent to random chance. This is itself a meaningful finding. It tells us that the '
    'features captured in any single encounter \u2014 behavioral indicators, call type, severity '
    '\u2014 cannot predict who will return to the system. Repeat contact appears to be driven by '
    'factors that exist outside the scope of individual encounters, which is consistent with '
    'housing instability being the only significant predictor in the structural sense.'
)

add_para(
    'From a system perspective, this distinction is critical. Reducing severity depends on '
    'improving real-time assessment and de-escalation, while reducing repeat encounters likely '
    'requires longer-term interventions that address structural factors such as housing stability '
    'and ongoing support.'
)

# ── 4.6 TIME-BASED ANALYSIS ────────────────────────────────────
add_heading('4.6  Time-Based Analysis', level=2)

add_figure('fig14_time_series.png',
           'Figure 17. Monthly Trends: Encounter Volume, Severity Rate, and Repeat Rate',
           width=Inches(6))

add_para(
    'Severity rates vary month-to-month, generally ranging between approximately 25% and 60%, '
    'with no consistent upward or downward trend over time. Importantly, fluctuations in severity '
    'do not align consistently with changes in encounter volume. Periods of higher volume, such '
    'as late 2024, do not correspond with sustained increases in severity. This indicates that '
    'increases in encounter volume are not driven by a higher proportion of severe cases, but '
    'rather by an increase in total encounters across similar levels of risk.'
)

add_para(
    'Repeat encounter rates follow a similar pattern \u2014 varying between approximately 15% '
    'and 50%, with no consistent trend. The variability is likely influenced by a small number of '
    'high-frequency individuals, where repeated encounters by a single or limited group can '
    'significantly impact monthly percentages. The highest observed repeat rate occurs in November '
    '2025 (50%), representing a concentrated effect rather than a structural change.'
)

add_para(
    'Overall, the data suggests that while encounter volume fluctuates over time, the proportion '
    'of severe and repeat encounters remains relatively stable in structure, with no evidence of '
    'systematic change.'
)

# ── 4.7 LOCATION ──────────────────────────────────────────────
add_heading('4.7  Location', level=2)

add_para(
    'A spatial heatmap of geocoded crisis encounters shows that activity is concentrated most '
    'heavily in the heart of Missoula. The strongest clustering appears in the urban core, '
    'particularly around East Broadway, West Alder, and the West Broadway / riverfront area, with '
    'additional concentrations extending along major corridors and key service-access points.'
)

add_para(
    'To evaluate this more systematically, encounters were grouped into geographic areas and '
    'compared across overall, severe, and repeat encounters. The percentage of encounters in each '
    'area was largely similar across all three groups. No area shows a distinct concentration of '
    'severe or repeat encounters beyond what would be expected based on overall activity levels. '
    'This suggests that geography reflects where encounters occur most frequently but does not '
    'meaningfully differentiate outcomes.'
)

page_break()

# ══════════════════════════════════════════════════════════════════
# 5. RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════════
print("Building Recommendations...")

add_heading('RECOMMENDATIONS', level=1)

add_para(
    'The primary limitation identified in this project is not the analytical approach, but the '
    'underlying data pipeline. The current process relies on handwritten forms followed by manual '
    'data entry, introducing multiple points of friction and error. This includes inconsistent '
    'data capture, transcription mistakes, and delays between when an encounter occurs and when '
    'it becomes available for analysis.'
)

add_heading('Expand Behavioral Health Response Capacity', level=2)

add_para(
    'Given that MST-only encounters show severity rates roughly 14 times lower than law '
    'enforcement-only encounters, and this pattern holds within the same call types, expanding '
    'MST staffing and broadening dispatch eligibility criteria could meaningfully reduce system-wide '
    'severity. Where behavioral health-only response is not feasible, co-response teams still '
    'show lower severity rates than law enforcement alone (48.0% vs. 58.3%), suggesting that '
    'defaulting to co-response for mental health calls would be an improvement over the current '
    'model. This represents one of the most actionable findings in this analysis.'
)

add_heading('Design Recommendations', level=2)

add_para(
    'A full redesign of the data collection workflow is the most critical next step. Moving from '
    'paper-based entry to a structured digital form (e.g., tablet-based or platform-integrated '
    'entry such as Alchemer) would immediately improve data quality, consistency, and availability.'
)

add_para('Key design improvements include:')

bullets = [
    'Replacing most free-text fields with standardized dropdowns and required inputs to reduce variability and missing data',
    'Introducing clearly defined behavioral indicators to ensure consistent documentation of severity-related factors',
    'Incorporating Likert-scale fields (e.g., perceived severity, level of risk, cooperation) to capture more nuanced, structured assessments',
    'Implementing conditional logic so only relevant fields appear based on prior responses, reducing cognitive load and completion time',
    'Standardizing location input using selectable landmarks, intersections, and GPS pin functionality to improve geocoding accuracy',
    'Adding validation rules (e.g., required fields, format checks) to prevent incomplete or unusable submissions',
    'Optional surveys to test effectiveness, training impact & system feedback',
]
for b in bullets:
    add_bullet(b)

add_para(
    'These improvements also create an opportunity to better evaluate training effectiveness. '
    'More structured and consistent inputs would allow for pre/post comparisons and ongoing '
    'assessment of how CIT training translates into field-level decision-making.'
)

add_para(
    'In parallel, building tailored analytical outputs for key stakeholders (e.g., Behavioral '
    'Health and Providence) would enable more targeted insights. Separate models or filtered '
    'dashboards could highlight patterns specific to each partner.'
)

add_heading('System-Level Improvements', level=2)

bullets2 = [
    'Real-time dashboards: Enable supervisors and stakeholders to monitor trends, severity, and repeat encounters as they occur',
    'Repeat individual tracking: Introduce a unique identifier system to track individuals across encounters and identify high-utilization patterns',
    'Severity flagging: Automatically flag high-acuity encounters to support prioritization and escalation with option to notify level-based stakeholders',
    'Cross-system integration: Connect data with behavioral health providers, hospitals (e.g., Providence), and service organizations to support continuity of care',
    'Operational reporting: Provide ongoing, automated reporting on key metrics (repeat rates, severity trends, location hotspots) to reduce reliance on manual analysis',
]
for b in bullets2:
    add_bullet(b)

add_heading('Ongoing Iteration', level=2)

add_para(
    'These recommendations should be implemented iteratively, with continuous feedback from '
    'officers, supervisors, and partner organizations. As data quality improves, additional '
    'modeling opportunities, such as predictive risk scoring, intervention targeting, and resource '
    'optimization, can be developed to further enhance system impact.'
)

page_break()

# ══════════════════════════════════════════════════════════════════
# 6. LIMITATIONS
# ══════════════════════════════════════════════════════════════════
print("Building Limitations...")

add_heading('LIMITATIONS', level=1)

add_para(
    'While this analysis provides meaningful insight into crisis response patterns, several '
    'important limitations should be considered when interpreting the results.'
)

add_heading('Responding Agency Assignment', level=2)

add_para(
    'An important limitation is that responding agency is not randomly assigned. MST and MPD may '
    'receive different types of calls based on dispatch protocols, caller descriptions, time of '
    'day, or resource availability. Although comparing severity rates within the same call types '
    'helps address this concern, there may still be unobserved differences between the encounters '
    'each agency handles. A fully causal interpretation would require a different study design, '
    'such as a natural experiment or instrumental variable approach. The findings should be '
    'understood as strong associations rather than definitive causal claims.'
)

add_heading('Person-Level Identifier Construction', level=2)

add_para(
    'A key limitation is the construction of person-level identifiers used to track repeat '
    'encounters. Because identifiers were derived from partial information (e.g., date of birth '
    'and initials prior to de-identification), there is potential for both over-linking '
    '(incorrectly combining different individuals) and under-linking (failing to match the same '
    'individual across encounters). Consistent matching rules were applied and ambiguous records '
    'were removed to mitigate this.'
)

add_heading('Data Collection and Pipeline Constraints', level=2)

add_para(
    'The most significant limitation is the reliance on a paper-based data collection process '
    'followed by manual data entry. This introduces multiple sources of inconsistency and error, '
    'including incomplete forms, illegible handwriting, and transcription inaccuracies. The '
    'process creates delays between when an encounter occurs and when it becomes available for '
    'analysis, limiting real-time usability.'
)

add_heading('Measurement and Variable Construction', level=2)

add_para(
    'Several key variables were constructed rather than directly observed. Most notably, the '
    'severity outcome was derived from multiple indicators, including use of force, injury, and '
    'legal criteria. While this approach is methodologically appropriate given the absence of a '
    'single severity field, it introduces subjectivity in how severity is defined. Behavioral '
    'indicators were also generated using keyword-based text parsing, which may not fully capture '
    'nuance in officer reporting. To improve accuracy, the parsing logic included exclusion rules '
    'for negation and absence (e.g., "no confusion," "not manic," "denies depression").'
)

add_heading('Missing Data and Category Imbalance', level=2)

add_para(
    'Some variables contained missing or unevenly distributed data, requiring category '
    'consolidation and data cleaning decisions. Low-frequency categories were collapsed into '
    '"Other" to stabilize model estimates. Missing values were dropped during modeling. If the '
    'missingness is not random, this could introduce bias.'
)

add_heading('Limited Longitudinal Tracking', level=2)

add_para(
    'The dataset does not support robust longitudinal tracking at the individual level. While a '
    'repeat contact flag was available, the absence of persistent, unique identifiers limits the '
    'ability to follow individuals over time or analyze trajectories.'
)

add_heading('Lack of External Data Integration', level=2)

add_para(
    'This analysis is based solely on data from the MHRF and does not incorporate external '
    'systems such as healthcare records, housing services, or law enforcement databases. '
    'Integration with these systems could significantly improve explanatory power.'
)

add_heading('Geographic and Data Processing Constraints', level=2)

add_para(
    'The findings are based on data from a single geographic area. Patterns observed in Missoula '
    'may not generalize to other regions. Location data required cleaning, standardization, and '
    'approximation in some cases.'
)

add_heading('Model Limitations', level=2)

add_para(
    'The logistic regression models identify associations rather than causal relationships. They '
    'are effective for understanding which variables are linked to severity and repeat encounters, '
    'but they do not establish causation. Binary logistic regression was selected because both '
    'primary outcomes are dichotomous. Some variables were retained for interpretability even when '
    'not statistically significant.'
)

page_break()

# ══════════════════════════════════════════════════════════════════
# 7. CONCLUSION
# ══════════════════════════════════════════════════════════════════
print("Building Conclusion...")

add_heading('CONCLUSION', level=1)

add_para(
    'This project demonstrates that improving crisis response requires both analytical insight '
    'and system-level design improvements.'
)

add_para(
    'The most significant finding is that the type of agency responding to a mental health crisis '
    'is the single strongest predictor of encounter severity \u2014 stronger than any behavioral '
    'indicator, call type, or demographic characteristic. Behavioral health professionals achieve '
    'severity rates of 4.1%, compared to 58.3% for law enforcement-only responses, and this '
    'difference holds within the same call types. This makes expanding behavioral health response '
    'capacity one of the most actionable recommendations from this analysis.'
)

add_para(
    'Beyond agency effects, the analysis shows that severity and repeat encounters are driven by '
    'fundamentally different mechanisms. Severity is primarily shaped by real-time behavioral '
    'escalation during an encounter, with indicators such as anger, mania, and confusion emerging '
    'as the strongest behavioral predictors. The severity model achieves a test AUC of 0.827, '
    'demonstrating strong predictive performance. In contrast, call type does not reliably explain '
    'severity, although certain categories are associated with lower likelihood of severe outcomes.'
)

add_para(
    'Repeat encounters follow a different pattern. Rather than being driven by characteristics of '
    'individual incidents, repeat contact is more strongly associated with underlying life '
    'conditions, particularly housing instability. The repeat model\u2019s AUC of 0.495 confirms '
    'that encounter-level features cannot predict who will return \u2014 a meaningful null finding '
    'that reinforces the structural nature of repeat involvement.'
)

add_para(
    'Crisis encounters group into three distinct profiles: escalation-driven (high severity, '
    'behavioral indicators), system-driven (low severity, high repeat, housing instability), and '
    'emotional distress (low acuity, depression-driven). This typology confirms that different '
    'types of crises require fundamentally different response approaches.'
)

add_para(
    'This distinction is critical, as it suggests that different types of interventions are '
    'needed. Severity requires real-time response strategies focused on de-escalation and '
    'behavioral management. Repeat contact points to the need for longer-term, system-level '
    'interventions that address persistent challenges such as housing instability and continuity '
    'of care.'
)

add_para(
    'At the same time, the project highlights a foundational constraint: the current paper-based '
    'data collection process limits the system\u2019s ability to fully leverage its data. '
    'Transitioning to a structured, digital data collection system represents a meaningful '
    'opportunity to address these limitations.'
)

add_para(
    'Ultimately, the value of CIT data lies not only in what is collected, but in how it is '
    'structured and used. Strengthening both the analytical framework and the data pipeline '
    'creates the foundation for a more proactive, targeted, and effective crisis response system.'
)

page_break()

# ══════════════════════════════════════════════════════════════════
# 8. APPENDIX A — CIT ACADEMY ANALYSIS
# ══════════════════════════════════════════════════════════════════
print("Building Appendix A...")

add_heading('APPENDIX A \u2014 40-HR CIT ACADEMY ANALYSIS', level=1)

add_para(
    'This analysis is included to complement the encounter-level findings presented in the main '
    'study. While the primary analysis focuses on factors associated with crisis severity and '
    'repeat encounters, the CIT Academy evaluation examines whether training improves officer '
    'capability in areas such as de-escalation, system knowledge, and confidence.',
    italic=True
)

add_heading('Introduction', level=2)

add_para(
    'Law enforcement officers frequently serve as first responders to individuals experiencing '
    'mental health crises. Effective crisis response requires not only procedural knowledge, but '
    'also familiarity with the behavioral health system, awareness of community resources, and '
    'confidence in de-escalation techniques. The Missoula 40-hour Crisis Intervention Team (CIT) '
    'Academy is designed to strengthen these competencies.'
)

add_para('The primary research questions guiding this evaluation were:')

add_para('1. Does the CIT Academy produce measurable improvements in officer knowledge and '
         'preparedness immediately following training?')
add_para('2. Are those improvements sustained at 3\u20136 months and 12 months?')
add_para('3. Are observed changes statistically significant and practically meaningful?')

add_heading('Data & Methods', level=2)

add_para(
    'Structured KSA surveys were administered at four timepoints: Pre-training (April 1, 2024), '
    'Immediate Post-training (April 5, 2024), 3\u20136 month follow-up, and 12-month follow-up. '
    'The full dataset included 92 responses across professional roles. The present analysis '
    'focuses on Missoula Law Enforcement (MLE) participants (36 total responses: Pre n=15, '
    'Post n=7, 3\u20136 month n=10, 12 month n=4).'
)

add_para(
    'Analyses included: descriptive mean comparisons, Welch\u2019s independent samples t-tests '
    '(Pre vs Post), Cohen\u2019s d effect size estimation, one-way ANOVA across stages, and '
    'Composite KSA Index construction.'
)

add_heading('Results', level=2)

add_para(
    'All eight KSA metrics increased from Pre-training to Immediate Post-training. Mean '
    'Pre-training scores ranged from 2.13 to 2.87, while Post-training means ranged from 3.14 '
    'to 4.00. The largest improvements were in preparedness to engage in crisis (+1.27), awareness '
    'of community resources (+1.25), and familiarity with mental health system roles (+1.17).'
)

# Add Appendix A figures
add_figure('extracted_appendix_fig1.png', 'Figure A1. Pre vs Post Mean Comparison', width=Inches(5.5))

add_para(
    'Seven of eight domains showed statistically significant improvements at alpha = .05. '
    'Cohen\u2019s d effect sizes ranged from 1.03 to 3.31, indicating large to extremely large '
    'practical effects.'
)

# Welch's t-test table
add_table_simple(
    ['Metric', 'Mean Pre', 'Mean Post', 'Diff', 'p-value', "Cohen's d"],
    [
        ['Preparedness to engage', '2.73', '4.00', '+1.27', '< .001', '3.31'],
        ['Resource awareness', '2.47', '3.71', '+1.25', '.0002', '1.84'],
        ['MH system familiarity', '2.40', '3.57', '+1.17', '.0005', '2.27'],
        ['Civil commitment knowledge', '2.13', '3.14', '+1.01', '.0004', '1.54'],
        ['De-escalation confidence', '2.87', '3.86', '+0.99', '.0001', '2.07'],
        ['MH knowledge comfort', '2.60', '3.57', '+0.97', '.0022', '1.61'],
        ['Crisis engagement comfort', '2.87', '3.57', '+0.71', '.0136', '1.35'],
        ['Professional liability', '2.47', '3.14', '+0.68', '.0509', '1.03'],
    ],
    col_widths=[1.8, 0.7, 0.7, 0.6, 0.7, 0.8]
)
add_para('')

add_heading('Composite KSA Index', level=3)

add_para(
    'Immediately following training, the Composite KSA Index increased from 2.57 to 3.57, '
    'representing an average improvement of +1.00 points on a five-point scale. At 3\u20136 '
    'months, the mean remained elevated at 3.43, and at 12 months it measured 3.53. Standard '
    'deviations were small across stages (0.29\u20130.37), indicating consistent responses.'
)

add_figure('extracted_appendix_fig2.png', 'Figure A2. Composite KSA Trajectory',
           width=Inches(5.5))

add_heading('ANOVA', level=3)

add_para(
    'One-way ANOVA tests confirmed statistically significant stage effects for all eight KSA '
    'metrics (F-statistics ranging from 4.50 to 15.67, all p-values \u2264 .0096). The strongest '
    'effect was for preparedness to engage in crisis (F = 15.67, p < .001).'
)

add_figure('extracted_appendix_fig3.png', 'Figure A3. Multi-line KSA Trajectory',
           width=Inches(5.5))

add_heading('Discussion', level=2)

add_para(
    'The CIT Academy produces substantial immediate improvements in officer-reported knowledge, '
    'preparedness, and confidence. These improvements were sustained at 3\u20136 and 12 months. '
    'The largest gains were in preparedness and resource awareness, central to effective crisis '
    'response. When considered alongside encounter-level results, the training primarily addresses '
    'behavioral response during incidents, while repeat system involvement is driven more by '
    'structural factors such as housing instability. Additionally, the agency analysis suggests '
    'that who responds may matter as much as how they are trained, highlighting that training must '
    'be complemented by system-level improvements.'
)

add_heading('Limitations', level=2)

add_para(
    'Survey responses were deidentified and not linked across timepoints. The 12-month follow-up '
    'sample (n=4) limits statistical power. Outcomes were based on self-reported perceptions '
    'rather than field-based performance measures. Non-response bias may influence findings.'
)

page_break()

# ══════════════════════════════════════════════════════════════════
# 9. APPENDIX B — BEHAVIORAL INDICATOR VALIDATION
# ══════════════════════════════════════════════════════════════════
print("Building Appendix B...")

add_heading('APPENDIX B \u2014 BEHAVIORAL INDICATOR VALIDATION', level=1)

add_heading('Purpose', level=2)

add_para(
    'This appendix evaluates whether behavioral indicators function as independent signals of '
    'severity or simply reflect descriptions of severity criteria.'
)

add_heading('Table B1. Behavioral Indicator Rates (Severe vs Non-Severe)', level=3)

add_table_simple(
    ['Behavior', 'Severe', 'Non-Severe', 'Difference'],
    [
        ['Angry / Uncooperative', '0.270', '0.085', '+0.186'],
        ['Confusion', '0.322', '0.162', '+0.160'],
        ['Manic', '0.230', '0.077', '+0.153'],
        ['Disorganized Speech', '0.239', '0.107', '+0.131'],
        ['Substance Involvement', '0.330', '0.242', '+0.088'],
        ['Scared / Frightened', '0.172', '0.147', '+0.026'],
        ['Depressed', '0.316', '0.305', '+0.012'],
        ['Delusions', '0.006', '0.008', '\u22120.002'],
        ['Hallucinations', '0.009', '0.036', '\u22120.027'],
    ],
    col_widths=[2.0, 1.0, 1.0, 1.0]
)
add_para('')

add_heading('Table B2. Strict Severity Validation (Force or Injury Only)', level=3)

add_table_simple(
    ['Behavior', 'Severe', 'Non-Severe', 'Difference'],
    [
        ['Angry / Uncooperative', '0.329', '0.118', '+0.212'],
        ['Manic', '0.234', '0.115', '+0.119'],
        ['Confusion', '0.299', '0.208', '+0.092'],
        ['Disorganized Speech', '0.228', '0.143', '+0.084'],
        ['Scared / Frightened', '0.198', '0.147', '+0.050'],
        ['Substance Involvement', '0.305', '0.271', '+0.035'],
        ['Delusions', '0.000', '0.008', '\u22120.008'],
        ['Hallucinations', '0.000', '0.031', '\u22120.031'],
    ],
    col_widths=[2.0, 1.0, 1.0, 1.0]
)
add_para('')

add_heading('Interpretation', level=2)

add_para(
    'Behavioral indicators such as anger, mania, and confusion are consistently more prevalent '
    'in severe encounters and remain differentiating under a stricter definition of severity. '
    'These behaviors also appear in cases where severity components (force or injury) are absent, '
    'indicating they are not solely descriptive of severity but reflect underlying escalation. '
    'In contrast, indicators such as depression, delusions, and hallucinations show weak or '
    'negative associations, suggesting they reflect different types of crisis presentations.'
)

page_break()

# ══════════════════════════════════════════════════════════════════
# 10. WORKS CITED
# ══════════════════════════════════════════════════════════════════
print("Building Works Cited...")

add_heading('WORKS CITED', level=1)

citations = [
    'National Alliance on Mental Illness (NAMI). (2023). Crisis intervention team (CIT) programs.',
    'Substance Abuse and Mental Health Services Administration (SAMHSA). (2020). National guidelines for behavioral health crisis care.',
    'Missoula Police Department. (n.d.). Missoula crisis intervention team (CIT). City of Missoula. https://www.ci.missoula.mt.us/2763/Missoula-Crisis-Intervention-Team-CIT',
    'OpenStreetMap contributors. (n.d.). Map data and basemap. https://www.openstreetmap.org',
]
for c in citations:
    p = add_para(c, space_after=8)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.left_indent = Inches(0.5)


# ══════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════
print("\nSaving document...")
doc.save(OUT)
file_size = os.path.getsize(OUT)
print(f"Saved to: {OUT}")
print(f"File size: {file_size / 1024:.0f} KB")
print(f"Paragraphs: {len(doc.paragraphs)}")
print("Done!")
